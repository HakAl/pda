import os
import logging
import time
from typing import List, Dict, NamedTuple, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
from contextlib import contextmanager
from typing import List, Tuple, Optional
from langchain_community.document_loaders import PyPDFLoader, TextLoader, CSVLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain.chains import HypotheticalDocumentEmbedder
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import OllamaLLM
from rank_bm25 import BM25Okapi
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed


logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class LoadResult(NamedTuple):
    """Data structure for returning loader results."""
    loaded_documents: List[Document]
    failed_files: List[Dict[str, str]]


try:
    from langchain_community.document_loaders import (
        Docx2txtLoader,
        CSVLoader
    )

    EXTRA_LOADERS_AVAILABLE = True
except ImportError:
    EXTRA_LOADERS_AVAILABLE = False
    print("⚠️  Additional document loaders not available. Install with: pip install docx2txt unstructured")


class ProgressEmbeddings(Embeddings):
    """Wrapper class that tracks embedding progress with tqdm."""

    def __init__(self, base_embeddings, total_chunks):
        self.base = base_embeddings
        self.pbar = tqdm(total=total_chunks, desc="🔢 Embedding chunks")
        self.embedded_count = 0

    def embed_documents(self, texts):
        result = self.base.embed_documents(texts)
        self.pbar.update(len(texts))
        self.embedded_count += len(texts)
        return result

    def embed_query(self, text):
        return self.base.embed_query(text)

    def close(self):
        self.pbar.close()

@contextmanager
def create_progress_embeddings(base_embeddings, total_chunks):
    """Context manager for progress-tracked embeddings.

    Usage:
        with progress_embeddings(embeddings, len(chunks)) as prog_emb:
            vector_store = Chroma.from_documents(
                documents=chunks,
                embedding=prog_emb,
                ...
            )
    """
    wrapper = ProgressEmbeddings(base_embeddings, total_chunks)
    try:
        yield wrapper
    finally:
        wrapper.close()


class DocumentProcessor:
    def __init__(self, persist_directory="./chroma_db", api_key=None):
        self.persist_directory = persist_directory
        self.api_key = api_key
        self.embeddings = self._setup_embeddings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=100,
            length_function=len,
        )
        self.bm25_index = None
        self.bm25_chunks = None
        self.supported_extensions = {
            '.pdf', '.txt',  '.docx', '.doc', '.csv'
        }

    def _setup_embeddings(self):
        """Setup embeddings with caching enabled"""
        try:
            return HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': False},
                cache_folder="./model_cache"
            )
        except ImportError:
            from langchain.embeddings import FakeEmbeddings
            return FakeEmbeddings(size=384)

    def load_documents(self, documents_folder: str = "./documents", retries: int = 2, delay: int = 1) -> LoadResult:
        """
        Load documents with robust parallel processing, retry logic, and failure tracking.

        Args:
            documents_folder: The folder to load files from.
            retries: The number of times to retry loading a failed document.
            delay: The initial delay in seconds between retries.

        Returns:
            A LoadResult object containing lists of loaded documents and failed files.
        """
        if not os.path.exists(documents_folder):
            logging.warning(f"Folder '{documents_folder}' not found.")
            return LoadResult([], [])

        file_paths = [
            os.path.join(documents_folder, f)
            for f in os.listdir(documents_folder)
            if any(f.lower().endswith(ext) for ext in self.supported_extensions)
        ]

        if not file_paths:
            logging.info(f"No supported files found in '{documents_folder}'.")
            logging.info(f"Supported formats: {', '.join(sorted(self.supported_extensions))}")
            return LoadResult([], [])

        loaded_docs: List[Document] = []
        failed_files: List[Dict[str, str]] = []
        max_workers = min(len(file_paths), os.cpu_count() or 4, 8)

        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_path = {
                executor.submit(self._load_single_document, fp, retries, delay): fp
                for fp in file_paths
            }

            with tqdm(total=len(file_paths), desc="📄 Loading files") as pbar:
                for future in as_completed(future_to_path):
                    file_path = future_to_path[future]
                    try:
                        docs = future.result()
                        if docs:
                            loaded_docs.extend(docs)
                    except Exception as e:
                        # NEW: Track failed files and the reason
                        failed_files.append({"path": file_path, "error": str(e)})
                        logging.error(f"Failed to process {os.path.basename(file_path)} after all retries: {e}")
                    finally:
                        pbar.update(1)

        return LoadResult(loaded_documents=loaded_docs, failed_files=failed_files)

    def _load_single_document(self, file_path: str, retries: int, delay: int) -> List[Document]:
        """
        Loads a single document with a retry mechanism.

        Args:
            file_path: Path to the file.
            retries: Number of retry attempts.
            delay: Initial delay between retries.

        Returns:
            A list of Document objects.

        Raises:
            Exception: If loading fails after all retry attempts.
        """
        last_exception = None
        for attempt in range(retries + 1):
            try:
                filename = os.path.basename(file_path)
                file_ext = os.path.splitext(file_path)[1].lower()

                loader_map = {
                    '.pdf': PyPDFLoader(file_path),
                    '.txt': TextLoader(file_path, encoding='utf-8'),
                    '.csv': CSVLoader(file_path),
                }

                if file_ext in ['.docx', '.doc']:
                    loaded_docs = self._load_word_simple(file_path, filename)
                elif file_ext in loader_map:
                    loader = loader_map[file_ext]
                    loaded_docs = loader.load()
                else:
                    return []  # Not a supported file, but not an error

                for doc in loaded_docs:
                    doc.metadata.update({
                        'source': filename,
                        'file_path': file_path,
                        'file_type': file_ext
                    })
                return loaded_docs  # Success

            except Exception as e:
                last_exception = e
                if attempt < retries:
                    time.sleep(delay * (2 ** attempt))  # Exponential backoff
                    logging.warning(f"Retrying ({attempt + 1}/{retries}) for {filename} due to error: {e}")
                else:
                    # Re-raise the last exception if all retries fail
                    raise Exception(f"Failed to load {filename}: {str(last_exception)}") from last_exception
        return []  # Should not be reached

    def _load_word_simple(self, file_path: str, filename: str) -> List[Document]:
        """Simple Word document loader using python-docx directly."""
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            raise ImportError(
                "python-docx is required for Word documents. Install with: pip install python-docx") from e

        try:
            doc = DocxDocument(file_path)
            full_text = [p.text for p in doc.paragraphs if p.text.strip()]

            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            content = "\n".join(full_text)
            if not content.strip():
                raise ValueError("No readable text found in Word document")

            return [Document(
                page_content=content,
                metadata={'source': filename, 'file_path': file_path, 'file_type': '.docx'}
            )]
        except Exception as e:
            # Propagate a specific error for better debugging
            raise Exception(f"Word document parsing error: {e}") from e


    def _build_bm25_index(self, chunks: List[Document]) -> Tuple[BM25Okapi, List[Document]]:
        """Build BM25 index with pre-tokenized corpus"""
        print("🔍 Building BM25 index...")
        # Tokenize once and cache
        tokenized = [doc.page_content.lower().split() for doc in chunks]
        return BM25Okapi(tokenized), chunks

    def _split_documents_batch(self, documents: List[Document]) -> List[Document]:
        return self.text_splitter.split_documents(documents)

    def process_documents(self, documents_folder: str = "./documents") -> Optional[Tuple]:
        """
        Loads, splits, and indexes documents from a folder, now with robust error handling.
        """
        logging.info("📂 Loading documents...")
        load_result = self.load_documents(documents_folder)

        if load_result.failed_files:
            logging.warning(f"⚠️ Encountered {len(load_result.failed_files)} loading errors:")
            for failed in load_result.failed_files:
                logging.warning(f"  - File: {os.path.basename(failed['path'])}, Error: {failed['error']}")

        if not load_result.loaded_documents:
            logging.error("❌ No documents were successfully loaded!")
            logging.info(f"   Please check the files in '{documents_folder}' and review any errors above.")
            return None

        documents = load_result.loaded_documents
        logging.info(f"✅ Successfully loaded {len(documents)} document pages.")

        logging.info("✂️  Splitting into chunks...")
        docs_to_split = []
        pre_chunked_docs = []

        for doc in documents:
            if doc.metadata.get('doc_type') in ['sheet_overview', 'column_descriptions', 'data_chunk']:
                pre_chunked_docs.append(doc)
            else:
                docs_to_split.append(doc)

        split_chunks = self._split_documents_batch(docs_to_split)
        all_final_chunks = split_chunks + pre_chunked_docs
        logging.info(f"✅ Created {len(all_final_chunks)} chunks (including pre-chunked data).")

        # Build BM25 index
        self.bm25_index, self.bm25_chunks = self._build_bm25_index(all_final_chunks)

        # Create vector store
        logging.info("🗄️  Creating vector database...")
        vector_store = self._create_vectorstore_batched(all_final_chunks)

        logging.info("✅ Vector database created successfully!")
        return vector_store, self.bm25_index, self.bm25_chunks

    def _create_vectorstore_batched(self, chunks: List[Document]) -> Chroma:
        """Create vector store with progress tracking via context manager."""

        with create_progress_embeddings(self.embeddings, len(chunks)) as prog_emb:
            vector_store = Chroma.from_documents(
                documents=chunks,
                embedding=prog_emb,
                persist_directory=self.persist_directory,
                collection_metadata={"hnsw:space": "cosine"}
            )

        return vector_store

    def load_existing_vectorstore(self) -> Tuple[Optional[Chroma], Optional[BM25Okapi], Optional[List[Document]]]:
        """Load existing vector store and rebuild BM25 index"""
        if not os.path.exists(self.persist_directory):
            print(f"ℹ️  No existing vector store found at '{self.persist_directory}'")
            return None, None, None

        try:
            print("📥 Loading existing vector store...")
            vs = Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embeddings
            )

            # Rebuild BM25 from stored chunks
            print("🔍 Rebuilding BM25 index...")
            raw = vs.get()

            if not raw["documents"]:
                print("⚠️  Vector store is empty")
                return vs, None, None

            chunks = [
                Document(page_content=doc, metadata=meta)
                for doc, meta in zip(raw["documents"], raw["metadatas"])
            ]

            self.bm25_index, self.bm25_chunks = self._build_bm25_index(chunks)
            print(f"✅ Loaded {len(chunks)} chunks from existing store")

            return vs, self.bm25_index, self.bm25_chunks

        except Exception as e:
            print(f"❌ Error loading vector store: {e}")
            return None, None, None

    def _create_hyde_embeddings(self):
        """Create HyDE embeddings wrapper"""
        llm = OllamaLLM(model="phi3:mini", temperature=0.1)
        return HypotheticalDocumentEmbedder.from_llm(
            llm=llm,
            base_embeddings=self.embeddings,
            prompt_key="web_search"
        )